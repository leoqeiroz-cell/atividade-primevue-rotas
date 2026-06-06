from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Relatorio_Cloud_Docker_PrimeShop.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def add_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(64, 64, 64)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Relatorio tecnico - Aplicacao containerizada em nuvem")
title_run.bold = True
title_run.font.name = "Arial"
title_run.font.size = Pt(18)
title_run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("PrimeShop | Docker, PaaS, escalabilidade e responsabilidade compartilhada")
subtitle_run.font.name = "Arial"
subtitle_run.font.size = Pt(10.5)
subtitle_run.italic = True

meta = doc.add_table(rows=3, cols=2)
meta.autofit = False
for row in meta.rows:
    row.cells[0].width = Inches(2.1)
    row.cells[1].width = Inches(5.0)
set_cell_text(meta.cell(0, 0), "Nome completo do aluno", True)
set_cell_text(meta.cell(0, 1), "[inserir nome completo]")
set_cell_text(meta.cell(1, 0), "Projeto", True)
set_cell_text(meta.cell(1, 1), "PrimeShop - vitrine web com area administrativa simulada")
set_cell_text(meta.cell(2, 0), "Modelo de servico", True)
set_cell_text(meta.cell(2, 1), "PaaS - Platform as a Service")
for row in meta.rows:
    set_cell_shading(row.cells[0], "D9EAF7")

add_heading(doc, "1. Apresentacao do projeto")
add_body(doc, "Este projeto simula a entrega de uma aplicacao web para um cliente que precisa de uma solucao segura, escalavel, hospedada em nuvem e baseada em containers. A aplicacao escolhida foi a PrimeShop, uma loja demonstrativa em Vue 3 com vitrine de produtos, carrinho, checkout protegido por login simulado e area administrativa restrita ao perfil de administrador.")
add_body(doc, "A proposta aproxima a atividade de um cenario real: a equipe entrega uma aplicacao pronta para producao, empacotada em Docker e servida por Nginx, em vez de depender de configuracoes manuais em cada ambiente.")

add_heading(doc, "2. Modelo de servico escolhido")
add_body(doc, "O modelo escolhido foi PaaS, usando como referencia servicos como Azure App Service for Containers, AWS Elastic Beanstalk com Docker ou Google Cloud Run. O PaaS reduz a carga operacional porque a equipe publica a imagem Docker e a plataforma gerenciada cuida de boa parte da execucao, disponibilidade, logs, HTTPS e escalabilidade.")
add_body(doc, "A mesma solucao poderia rodar em IaaS, como uma maquina virtual EC2, mas isso exigiria maior responsabilidade sobre sistema operacional, patches e configuracao do servidor. Para este projeto, PaaS e a opcao mais equilibrada.")

add_heading(doc, "3. Componentes utilizados")
add_bullets(doc, [
    "Vue 3, Vite, PrimeVue e Tailwind CSS para a aplicacao web.",
    "Docker para empacotar a aplicacao em uma imagem portavel.",
    "Node.js 20 Alpine como imagem oficial da etapa de build.",
    "Nginx 1.27 Alpine como imagem oficial da etapa de producao.",
    "Docker Compose para simular porta, rede isolada e volumes persistentes.",
    "Volumes para cache e logs do Nginx, ajudando em diagnostico e auditoria."
])

add_heading(doc, "4. Preparacao do ambiente com Docker")
add_body(doc, "O Dockerfile usa uma estrategia multi-stage. A primeira etapa instala dependencias e gera o build da aplicacao. A segunda etapa copia apenas os arquivos finais para o Nginx, reduzindo o tamanho da imagem e evitando levar dependencias de desenvolvimento para producao.")
add_code(doc, "Porta simulada: localhost:8080 -> container:80\nRede Docker: primeshop-net\nVolumes: primeshop-nginx-cache e primeshop-nginx-logs")

add_heading(doc, "5. Diagrama simplificado da arquitetura")
diagram = doc.add_table(rows=8, cols=1)
steps = [
    "Usuario acessa a aplicacao por HTTPS",
    "Plataforma PaaS recebe a requisicao",
    "Servico em nuvem executa a imagem Docker",
    "Container PrimeShop responde pela porta 80",
    "Nginx entrega os arquivos estaticos do Vue/Vite",
    "Rede Docker isola a aplicacao",
    "Volume de cache melhora a operacao do Nginx",
    "Volume de logs apoia monitoramento e auditoria",
]
for i, step in enumerate(steps):
    set_cell_text(diagram.cell(i, 0), step, i == 0)
    set_cell_shading(diagram.cell(i, 0), "EAF3F8" if i % 2 == 0 else "FFFFFF")

add_heading(doc, "6. Simulacao de deploy")
add_body(doc, "A estrategia recomendada e automatizada com CI/CD. A cada alteracao no repositorio, o pipeline instalaria dependencias, executaria validacoes, geraria o build, construiria a imagem Docker, publicaria em um registry e atualizaria o servico na nuvem.")
add_body(doc, "Em uma implantacao na Azure App Service for Containers, por exemplo, a plataforma buscaria a imagem no Azure Container Registry e executaria o container com HTTPS, variaveis de ambiente, monitoramento e possibilidade de aumento de instancias. Para reduzir riscos, o deploy pode ser escalonado: homologacao primeiro, producao depois.")

add_heading(doc, "7. Escalabilidade e elasticidade")
add_body(doc, "A escalabilidade acontece porque o container nao guarda estado interno relevante. Assim, a plataforma pode executar varias instancias identicas atras de um balanceador de carga. Se a quantidade de acessos aumentar, novas copias do container podem ser criadas.")
add_body(doc, "A elasticidade complementa esse comportamento: em horarios de pico, a nuvem aumenta recursos ou instancias; quando a demanda diminui, reduz novamente. Isso mantem disponibilidade sem desperdicar custo com capacidade ociosa.")

add_heading(doc, "8. Seguranca e responsabilidade compartilhada")
add_body(doc, "Na nuvem, a seguranca e compartilhada. O provedor cuida da infraestrutura fisica, rede base, datacenter e parte da plataforma. A equipe do projeto cuida do codigo, imagem Docker, configuracoes, credenciais, permissoes e atualizacoes.")
add_bullets(doc, [
    "Uso de imagens oficiais e pequenas.",
    "Separacao entre build e execucao.",
    "Publicacao apenas da porta necessaria.",
    "Healthcheck para monitoramento do container.",
    "Logs persistidos para diagnostico.",
    "Controle de acesso simulado na aplicacao."
])

add_heading(doc, "9. Beneficios, desafios e conclusao")
add_body(doc, "Os beneficios principais sao portabilidade, padronizacao, facilidade de deploy e melhor aproveitamento dos recursos da nuvem. A equipe consegue entregar a mesma imagem em ambiente local, homologacao ou producao.")
add_body(doc, "Os desafios envolvem manter imagens atualizadas, proteger segredos, observar logs, configurar limites de recursos e lembrar que containerizar nao torna a aplicacao automaticamente segura. A solucao atende ao enunciado porque apresenta Dockerfile funcional, porta de acesso, rede, volume persistente, deploy em nuvem e aplicacao dos conceitos de escalabilidade, elasticidade e responsabilidade compartilhada.")

add_heading(doc, "10. Execucao local")
add_code(doc, "docker compose up --build\nAcesso: http://localhost:8080\ndocker compose down")

doc.save(OUTPUT)
print(OUTPUT)
